from .base import BaseLoader
from docx import Document
from typing import Dict, Any
import os

class DocxLoader(BaseLoader):
    def load(self, file_path: str, **kwargs) -> str:
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text

    def load_metadata(self, file_path: str) -> Dict[str, Any]:
        doc = Document(file_path)
        return {
            "core_properties": doc.core_properties.title,
            "author": doc.core_properties.author
        }
