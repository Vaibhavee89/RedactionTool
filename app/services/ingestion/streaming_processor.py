import os
from typing import Iterator, Dict, Any, Callable, Optional, List
from pathlib import Path
from .pdf_loader import PDFLoader
from .docx_loader import DocxLoader
from .text_loader import TextLoader
from app.services.pii.detector_engine import DetectorEngine
from app.services.redaction.redactor import Redactor

class StreamingProcessor:
    """
    Streaming processor for chunked processing of large files.
    Processes files in chunks to avoid memory issues with very large documents.
    """

    def __init__(self, chunk_size: int = 10000):
        """
        Initialize streaming processor.

        Args:
            chunk_size: Number of characters per chunk (default: 10000)
        """
        self.chunk_size = chunk_size
        self.detector = DetectorEngine()
        self.redactor = Redactor()

    def process_large_text_file(
        self,
        file_path: str,
        output_path: str,
        overlap: int = 500,
        progress_callback: Optional[Callable[[float], None]] = None
    ) -> Dict[str, Any]:
        """
        Process a large text file in chunks.

        Args:
            file_path: Path to input file
            output_path: Path to output file
            overlap: Number of characters to overlap between chunks (to avoid missing PII at boundaries)
            progress_callback: Optional callback(progress) for tracking

        Returns:
            Dictionary with processing statistics
        """
        ext = Path(file_path).suffix.lower()

        if ext == '.txt':
            return self._process_txt_streaming(file_path, output_path, overlap, progress_callback)
        elif ext == '.pdf':
            return self._process_pdf_streaming(file_path, output_path, progress_callback)
        elif ext == '.docx':
            return self._process_docx_streaming(file_path, output_path, overlap, progress_callback)
        else:
            raise ValueError(f"Unsupported file type for streaming: {ext}")

    def _process_txt_streaming(
        self,
        file_path: str,
        output_path: str,
        overlap: int,
        progress_callback: Optional[Callable[[float], None]]
    ) -> Dict[str, Any]:
        """Stream process a plain text file."""
        file_size = os.path.getsize(file_path)
        bytes_processed = 0
        total_pii_found = 0

        with open(file_path, 'r', encoding='utf-8', errors='ignore') as infile, \
             open(output_path, 'w', encoding='utf-8') as outfile:

            buffer = ""
            for chunk in self._read_in_chunks(infile, self.chunk_size):
                buffer += chunk

                # Process buffer when it reaches chunk size
                if len(buffer) >= self.chunk_size:
                    # Process chunk
                    findings = self.detector.detect(buffer)
                    policy = {f['entity_type']: 'block' for f in findings}
                    redacted = self.redactor.redact_text(buffer, findings, policy)

                    # Write all but the overlap
                    write_amount = len(redacted) - overlap
                    if write_amount > 0:
                        outfile.write(redacted[:write_amount])

                    # Keep overlap for next iteration
                    buffer = redacted[write_amount:]

                    total_pii_found += len(findings)

                    # Update progress
                    bytes_processed += len(chunk.encode('utf-8'))
                    if progress_callback:
                        progress_callback(bytes_processed / file_size)

            # Process remaining buffer
            if buffer:
                findings = self.detector.detect(buffer)
                policy = {f['entity_type']: 'block' for f in findings}
                redacted = self.redactor.redact_text(buffer, findings, policy)
                outfile.write(redacted)
                total_pii_found += len(findings)

        if progress_callback:
            progress_callback(1.0)

        return {
            'file': file_path,
            'output': output_path,
            'total_pii_found': total_pii_found,
            'file_size_bytes': file_size
        }

    def _process_pdf_streaming(
        self,
        file_path: str,
        output_path: str,
        progress_callback: Optional[Callable[[float], None]]
    ) -> Dict[str, Any]:
        """Stream process a PDF file page by page."""
        import pdfplumber

        total_pii_found = 0
        processed_pages = 0

        with pdfplumber.open(file_path) as pdf, \
             open(output_path, 'w', encoding='utf-8') as outfile:

            total_pages = len(pdf.pages)

            for i, page in enumerate(pdf.pages):
                # Extract page text
                text = page.extract_text() or ""

                # Detect and redact
                findings = self.detector.detect(text)
                policy = {f['entity_type']: 'block' for f in findings}
                redacted = self.redactor.redact_text(text, findings, policy)

                # Write to output
                outfile.write(f"\n--- Page {i+1} ---\n")
                outfile.write(redacted)

                total_pii_found += len(findings)
                processed_pages += 1

                # Update progress
                if progress_callback:
                    progress_callback((i + 1) / total_pages)

        if progress_callback:
            progress_callback(1.0)

        return {
            'file': file_path,
            'output': output_path,
            'total_pii_found': total_pii_found,
            'total_pages': processed_pages
        }

    def _process_docx_streaming(
        self,
        file_path: str,
        output_path: str,
        overlap: int,
        progress_callback: Optional[Callable[[float], None]]
    ) -> Dict[str, Any]:
        """Stream process a DOCX file paragraph by paragraph."""
        from docx import Document

        doc = Document(file_path)
        total_paragraphs = len(doc.paragraphs)
        total_pii_found = 0

        with open(output_path, 'w', encoding='utf-8') as outfile:
            buffer = ""

            for i, paragraph in enumerate(doc.paragraphs):
                buffer += paragraph.text + "\n"

                # Process buffer when it reaches chunk size
                if len(buffer) >= self.chunk_size:
                    findings = self.detector.detect(buffer)
                    policy = {f['entity_type']: 'block' for f in findings}
                    redacted = self.redactor.redact_text(buffer, findings, policy)

                    # Write all but the overlap
                    write_amount = len(redacted) - overlap
                    if write_amount > 0:
                        outfile.write(redacted[:write_amount])

                    buffer = redacted[write_amount:]
                    total_pii_found += len(findings)

                # Update progress
                if progress_callback:
                    progress_callback((i + 1) / total_paragraphs)

            # Process remaining buffer
            if buffer:
                findings = self.detector.detect(buffer)
                policy = {f['entity_type']: 'block' for f in findings}
                redacted = self.redactor.redact_text(buffer, findings, policy)
                outfile.write(redacted)
                total_pii_found += len(findings)

        if progress_callback:
            progress_callback(1.0)

        return {
            'file': file_path,
            'output': output_path,
            'total_pii_found': total_pii_found,
            'total_paragraphs': total_paragraphs
        }

    def _read_in_chunks(self, file_obj, chunk_size: int) -> Iterator[str]:
        """Generator to read file in chunks."""
        while True:
            data = file_obj.read(chunk_size)
            if not data:
                break
            yield data

    def estimate_processing_time(self, file_path: str) -> Dict[str, Any]:
        """
        Estimate processing time based on file size.

        Returns:
            Dictionary with file stats and estimated time
        """
        file_size = os.path.getsize(file_path)
        ext = Path(file_path).suffix.lower()

        # Rough estimates (chars per second)
        processing_speed = {
            '.txt': 50000,  # Fast
            '.pdf': 20000,  # Moderate (depends on OCR)
            '.docx': 30000  # Moderate
        }

        speed = processing_speed.get(ext, 20000)
        estimated_chars = file_size  # Rough approximation
        estimated_seconds = estimated_chars / speed

        return {
            'file_size_bytes': file_size,
            'file_type': ext,
            'estimated_seconds': estimated_seconds,
            'estimated_minutes': estimated_seconds / 60,
            'recommended_chunk_size': self.chunk_size
        }
